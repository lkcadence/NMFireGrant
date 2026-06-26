"""Generate Prior-Year Prefill Tester Summary Word document."""

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex):
  shading = OxmlElement('w:shd')
  shading.set(qn('w:fill'), fill_hex)
  cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
  doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
  p = doc.add_paragraph()
  run = p.add_run(text)
  if bold:
    run.bold = True
  return p


def add_bullets(doc, items):
  for item in items:
    doc.add_paragraph(item, style='List Bullet')


def add_term_definitions(doc, terms):
  for term, definition in terms:
    p = doc.add_paragraph()
    term_run = p.add_run(term)
    term_run.bold = True
    p.add_run(f' — {definition}')


def add_section_table(doc, headers, rows):
  table = doc.add_table(rows=1, cols=len(headers))
  table.style = 'Table Grid'
  hdr = table.rows[0].cells
  for i, h in enumerate(headers):
    hdr[i].text = h
    set_cell_shading(hdr[i], 'D9E2F3')
    for p in hdr[i].paragraphs:
      for run in p.runs:
        run.bold = True
  for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
      cells[i].text = val
  doc.add_paragraph('')


def build_document():
  doc = Document()
  style = doc.styles['Normal']
  style.font.name = 'Calibri'
  style.font.size = Pt(11)

  title = doc.add_paragraph()
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  t_run = title.add_run('NMSFM Fire Grant Application\nPrior-Year Prefill Summary')
  t_run.bold = True
  t_run.font.size = Pt(18)

  sub = doc.add_paragraph()
  sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
  sub_run = sub.add_run(f'For QA / UAT Testers — {date.today().strftime("%B %d, %Y")}')
  sub_run.font.size = Pt(12)

  doc.add_paragraph('')

  add_heading(doc, 'Purpose', 2)
  add_para(
    doc,
    'This document describes, for each application section, which fields are '
    'automatically loaded from a prior fiscal year (FY) application when a '
    'department starts a new FY application. Use it to plan test cases and '
    'confirm expected behavior.'
  )

  add_heading(doc, 'Terms', 2)
  add_para(
    doc,
    'This document uses a few technical terms. Definitions below are written '
    'for testers, not developers.'
  )
  add_term_definitions(doc, [
    (
      'Scalar field',
      'A single-value field on the form — for example, a text box, number '
      'field, or Yes/No radio button. Scalars are not part of a grid/table.'
    ),
    (
      'Grid',
      'A table on the form where the user can add, edit, or remove multiple '
      'rows (e.g., Aid Districts, Apparatus equipment, Water Sources). Each row '
      'has its own set of column values.'
    ),
    (
      'Gate',
      'A Yes/No question that controls whether a related section of the form '
      'is shown or required. Examples: "Is Apparatus Part of Project?" and '
      '"Aid Agreements?" When the gate is Yes, the detail fields and/or grid '
      'are visible; when No, that section is typically hidden.'
    ),
    (
      'Part-of-project gate',
      'A specific gate on Apparatus, Communication Equipment, or PPE that '
      'asks whether that category is part of the grant project. If prior-year '
      'data includes grid rows but no saved Yes/No answer, the application '
      'may default the gate to Yes so the prefilled rows remain visible.'
    ),
    (
      'Prefill',
      'Automatically loading field values from a prior FY application into '
      'the current FY form when the user opens a section. Prefill is display '
      'only until the user saves.'
    ),
    (
      'Prior FY / fiscal year (FY)',
      'An earlier grant application year for the same department (e.g., FY2026 '
      'data used when starting FY2027).'
    ),
    (
      'Walk-back',
      'The process of searching backward through prior FY applications to find '
      'the nearest year that has saved data for a given section. Empty years '
      'in between are skipped.'
    ),
    (
      'Remap',
      'When grid rows are copied from a prior FY, each row is assigned new '
      'internal IDs and linked to the current FY application before display. '
      'This prevents prior-year records from being overwritten when the user '
      'saves.'
    ),
    (
      'Partial prefill',
      'Only some fields in a section are loaded from prior FY; other fields '
      'are left blank for the user to enter fresh data for the new FY.'
    ),
    (
      'Full prefill',
      'All user-entered fields in a section (scalars and grid) are loaded '
      'from prior FY. Water Availability is the only section with full prefill.'
    ),
    (
      'UI-only',
      'Data shown on screen from prior FY that has not yet been saved to the '
      'current FY application in the database. Saving or navigating away '
      'writes it to the current application.'
    ),
    (
      'Read-only mode',
      'The application is locked for editing (e.g., after submission or when '
      'viewing as admin). Fields may display prefilled or saved values but '
      'cannot be changed.'
    ),
  ])

  add_heading(doc, 'How Prior-Year Prefill Works', 2)
  add_bullets(doc, [
    'Prefill runs when the user opens a section and no saved data exists for '
    'the current FY.',
    'The system walks back to the nearest prior FY that has data for that '
    'section (skipping empty intermediate years).',
    'Prefilled data appears in the form only — nothing is written to the '
    'database until the user clicks Save or navigates away (which triggers save).',
    'Grid rows from prior FY are remapped to the current application (new IDs) '
    'before display.',
    'Uploaded documents are never copied from prior FY applications.',
    'When a section loads prior-year data, a message is shown (info banner on '
    'General Information; plain text on grid sections).',
  ])

  add_heading(doc, 'Sections With No Prior-Year Prefill', 2)
  add_bullets(doc, [
    'Instructions',
    'Budget Information',
    'Response History',
    'Training',
    'PPE',
    'Equipment Needs',
    'Grant Funding Justification',
    'Project Budget Sheet',
    'Signatures and Supporting Docs',
    'Project Description (header on all pages — current FY only)',
  ])

  add_heading(doc, 'Section-by-Section Detail', 1)

  # General Information
  add_heading(doc, '1. General Information', 2)
  add_para(doc, 'Prior-year prefill: YES (partial)', bold=True)
  add_para(doc, 'Trigger: No current-FY General Information record saved.')
  add_heading(doc, 'Fields prefilled from prior FY', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Fire Chief Name', ''],
    ['Phone', ''],
    ['Email', ''],
    ['Grant Source (Individual / County-Wide)', ''],
    ['County departments compliant', 'When County-Wide is selected'],
    ['City/Municipality vs County', ''],
    ['Department type (Career / Volunteer / Combined)', ''],
    ['Admin department checkbox', ''],
    ['Total firefighters', ''],
    ['FFI firefighters', ''],
    ['FFII firefighters', ''],
    ['FD member Yes/No', ''],
  ])
  add_heading(doc, 'Fields NOT prefilled', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['NERIS ID', 'Uses master FDID list; prior-year value is not copied'],
    ['Department name, address, city, state, zip, county', 'From current department record'],
    ['ISO rating, stations, admin buildings', 'From department UDFs'],
    ['Community type (Urban / Rural / Suburban)', 'Left blank'],
    ['Mailing address overrides', 'Left blank'],
    ['Person completing application', 'Left blank'],
  ])
  add_para(doc, 'Banner: Info alert naming the source FY.')

  # Community Information
  add_heading(doc, '2. Community Information', 2)
  add_para(doc, 'Prior-year prefill: YES (partial)', bold=True)
  add_heading(doc, 'Fields prefilled from prior FY', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Community Protected', ''],
    ['Aid Agreements Yes/No', 'Defaults to Yes if prior value unset and aid-district grid has rows'],
    ['Aid Districts grid', 'Number, Aid District name'],
  ])
  add_heading(doc, 'Fields NOT prefilled', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Number of homes protected', 'User must enter each FY'],
    ['Number of commercial properties', 'User must enter each FY'],
    ['Permanent resident population', 'User must enter each FY'],
  ])
  add_para(
    doc,
    'Note: Save/Next requires homes, commercial, and population even after '
    'grid prefill.'
  )
  add_para(doc, 'Banner: "Information Loaded from Previous Application".')

  # Water Availability
  add_heading(doc, '3. Water Availability', 2)
  add_para(doc, 'Prior-year prefill: YES (full section)', bold=True)
  add_heading(doc, 'Fields prefilled from prior FY', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Community hydrant system Yes/No', ''],
    ['Total available water capacity', ''],
    ['Water on wheels capacity', ''],
    ['Station water capacity', ''],
    ['Water storage tank at station Yes/No', ''],
    ['Additional water sources grid', 'Number, Water Source, Capacity'],
  ])
  add_para(doc, 'This is the only section where all scalar fields and the grid are prefilled.')
  add_para(doc, 'Banner: "Information Loaded from Previous Application".')

  # Apparatus
  add_heading(doc, '4. Apparatus', 2)
  add_para(doc, 'Prior-year prefill: YES (partial)', bold=True)
  add_heading(doc, 'Fields prefilled from prior FY', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Is Apparatus Part of Project Yes/No', 'Defaults to Yes if unset and equipment grid has rows'],
    ['Apparatus equipment grid', 'Number, Name, Vehicle Type, Year, VIN, Capacity, GPM, Test Date, Pass, Comments'],
  ])
  add_heading(doc, 'Fields NOT prefilled', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Pump tests conducted Yes/No', ''],
    ['Explain no pump tests', ''],
    ['Hose tests conducted Yes/No', ''],
    ['Explain no hose tests', ''],
    ['Uploaded apparatus documents', 'Current FY only'],
  ])
  add_para(
    doc,
    'If gate auto-selects Yes with a prefilled grid, user must still complete '
    'pump/hose questions before Save succeeds. Selecting No hides the section '
    'but preserves grid data.'
  )
  add_para(doc, 'Banner: "Information Loaded from Previous Application".')

  # Communication Equipment
  add_heading(doc, '5. Communication Equipment', 2)
  add_para(doc, 'Prior-year prefill: YES (partial)', bold=True)
  add_heading(doc, 'Fields prefilled from prior FY', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Is Communication Part of Project Yes/No', 'Defaults to Yes if unset and equipment grid has rows'],
    ['Communication equipment grid', 'Number, Equipment, Quantity'],
  ])
  add_heading(doc, 'Fields NOT prefilled', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Handheld radios, base stations, mobile radios', ''],
    ['Apparatus without radio Yes/No', ''],
    ['Law enforcement interoperability', ''],
    ['Emergency medical interoperability', ''],
    ['Other fire dept interoperability', ''],
    ['Other agency interoperability + description', ''],
    ['Areas not covered by repeater + description', ''],
    ['Admin comments', ''],
    ['Uploaded communication documents', 'Current FY only'],
  ])
  add_para(
    doc,
    'Largest gap: grid may prefill with gate on Yes, but many scalar fields '
    'remain blank until the user completes them.'
  )
  add_para(doc, 'Banner: "Information Loaded from Previous Application".')

  # Hazards/Threats
  add_heading(doc, '6. Hazards/Threats', 2)
  add_para(doc, 'Prior-year prefill: YES (grid only)', bold=True)
  add_heading(doc, 'Fields prefilled from prior FY', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Hazards/threats grid', 'Number, Hazard Type, Hazard Detail'],
  ])
  add_heading(doc, 'Fields NOT prefilled', 3)
  add_section_table(doc, ['Field', 'Notes'], [
    ['Admin comments', 'Internal use only'],
  ])
  add_para(doc, 'Banner: "Information Loaded from Previous Application".')

  # PPE note
  add_heading(doc, '7. PPE (no prior-year prefill)', 2)
  add_para(
    doc,
    'PPE does not load data from prior FY. When saved current-FY data exists, '
    'part-of-project gates may default to Yes if the respective grid has rows. '
    'Selecting No preserves grid data (data is not wiped).'
  )

  add_heading(doc, 'Quick Reference Matrix', 1)
  add_section_table(doc, ['Section', 'Prefill?', 'What is prefilled'], [
    ['Instructions', 'No', '—'],
    ['General Information', 'Partial', 'Contact, grant source, dept type, firefighter counts'],
    ['Budget Information', 'No', '—'],
    ['Community Information', 'Partial', 'Community name, aid gate, aid districts grid'],
    ['Response History', 'No', '—'],
    ['Water Availability', 'Full', 'All scalars + water sources grid'],
    ['Training', 'No', '—'],
    ['Apparatus', 'Partial', 'Part-of-project gate, apparatus grid'],
    ['Communication Equipment', 'Partial', 'Part-of-project gate, equipment grid'],
    ['PPE', 'No', '—'],
    ['Hazards/Threats', 'Partial', 'Hazards grid only'],
    ['Equipment Needs', 'No', '—'],
    ['Grant Funding Justification', 'No', '—'],
    ['Project Budget Sheet', 'No', '—'],
    ['Signatures and Supporting Docs', 'No', '—'],
  ])

  add_heading(doc, 'Suggested Test Scenarios', 2)
  add_bullets(doc, [
    'New FY application with prior FY data: open each section above and confirm '
    'only the listed fields appear.',
    'Save and reload: prefilled data should persist after Save, Next, Previous, '
    'and sidebar navigation.',
    'Missing intermediate year: if FY N-1 is empty but FY N-2 has data, confirm '
    'walk-back uses FY N-2.',
    'Part-of-project gates: sections with grids should default to Yes when rows '
    'exist; verify No preserves data.',
    'Community Information: confirm homes, commercial, and population stay blank '
    'and are required on save.',
    'Documents: confirm no prior-year uploads appear on Apparatus, Communication, '
    'Training, or PPE.',
    'Read-only mode: prefilled fields should not be editable when application is '
    'read-only.',
  ])

  add_heading(doc, 'Document Control', 2)
  add_para(doc, 'Source: NMSFM Fire Grant Application codebase (CVE Fixes branch).')
  add_para(doc, 'Related developer docs: docs/planning/prior-year-prefill-save-plan.md')

  return doc


if __name__ == '__main__':
  output = (
    r'i:\NMSFM REPOS\Fire Grant App - CVE Fixes\NMSFM_FGF_CVE\NMSFMFireGrantWF'
    r'\docs\Prior-Year-Prefill-Tester-Summary.docx'
  )
  build_document().save(output)
  print(f'Created: {output}')
