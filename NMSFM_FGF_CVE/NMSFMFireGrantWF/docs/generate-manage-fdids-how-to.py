"""Generate Manage FDIDs / Related Addresses Admin How-To Word document."""

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_update_fields_on_open(doc):
  """Prompt Word to refresh fields (including TOC) when the document opens."""
  settings = doc.settings.element
  existing = settings.find(qn('w:updateFields'))
  if existing is None:
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)
  else:
    existing.set(qn('w:val'), 'true')


def add_table_of_contents(doc, levels='2-3'):
  """
  Insert a Word TOC field for Heading 2–3 (main sections and subsections).
  Populates when opened in Word (see set_update_fields_on_open).
  """
  toc_title = doc.add_paragraph()
  toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  toc_run = toc_title.add_run('Table of Contents')
  toc_run.bold = True
  toc_run.font.size = Pt(14)

  paragraph = doc.add_paragraph()
  run = paragraph.add_run()

  fld_begin = OxmlElement('w:fldChar')
  fld_begin.set(qn('w:fldCharType'), 'begin')

  instr = OxmlElement('w:instrText')
  instr.set(qn('xml:space'), 'preserve')
  instr.text = f' TOC \\o "{levels}" \\h \\z \\u '

  fld_sep = OxmlElement('w:fldChar')
  fld_sep.set(qn('w:fldCharType'), 'separate')

  run._r.append(fld_begin)
  run._r.append(instr)
  run._r.append(fld_sep)

  paragraph.add_run(
    'Open in Microsoft Word to populate the table of contents.'
  )

  run_end = paragraph.add_run()
  fld_end = OxmlElement('w:fldChar')
  fld_end.set(qn('w:fldCharType'), 'end')
  run_end._r.append(fld_end)

  doc.add_paragraph('')
  page_break = doc.add_paragraph()
  page_break.add_run().add_break(WD_BREAK.PAGE)


def set_cell_shading(cell, fill_hex):
  shading = OxmlElement('w:shd')
  shading.set(qn('w:fill'), fill_hex)
  cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
  doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
  p = doc.add_paragraph()
  run = p.add_run(text)
  if bold:
    run.bold = True
  if italic:
    run.italic = True
  return p


def add_bullets(doc, items):
  for item in items:
    doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc, items):
  for item in items:
    doc.add_paragraph(item, style='List Number')


def add_section_table(doc, headers, rows):
  table = doc.add_table(rows=1, cols=len(headers))
  table.style = 'Table Grid'
  hdr = table.rows[0].cells
  for i, h in enumerate(headers):
    hdr[i].text = h
    set_cell_shading(hdr[i], 'D9E2F3')
    for p in hdr[i].paragraphs:
      for run in p.runs:
        run.bold = True
  for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
      cells[i].text = val
  doc.add_paragraph('')


def build_document():
  doc = Document()
  style = doc.styles['Normal']
  style.font.name = 'Calibri'
  style.font.size = Pt(11)

  title = doc.add_paragraph()
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  t_run = title.add_run(
    'NMSFM Fire Grant Application\n'
    'Manage FDIDs / Related Addresses\nAdmin How-To'
  )
  t_run.bold = True
  t_run.font.size = Pt(18)

  sub = doc.add_paragraph()
  sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
  sub_run = sub.add_run(
    f'For Web Administrators — {date.today().strftime("%B %d, %Y")}'
  )
  sub_run.font.size = Pt(12)

  doc.add_paragraph('')
  add_para(doc, 'Project: NMSFM Fire Grant Web Application')
  add_para(doc, 'Document version: 1.0')
  add_para(doc, 'Audience: Fire Grant web administrators')
  add_para(
    doc,
    'Page: Admin → Manage FDIDs / Related Addresses (/Admin/ManageFDIDs)'
  )
  add_para(
    doc,
    'Source markdown: docs/how-to/manage-fdids-related-addresses-how-to.md'
  )
  doc.add_paragraph('')

  add_table_of_contents(doc)

  add_heading(doc, '1. What this page is for', 2)
  add_para(
    doc,
    'The Manage FDIDs / Related Addresses page maintains the master list of '
    'fire department NERIS IDs (stored as FDID in the database) and their '
    'department names. Each row in this list is used across the Fire Grant '
    'application—for example:'
  )
  add_bullets(doc, [
    'General Information — NERIS ID prefill and department lookup',
    'Edit User — department dropdown labels',
    'Award, denial, and print reports — department name and address on documents',
  ])
  add_para(
    doc,
    'After the NERIS migration, department names in this master list may no '
    'longer match the Codepal fire department address records. This page lets you:'
  )
  add_numbered(doc, [
    'Add or update NERIS IDs and department names in the master list.',
    'Link an existing Codepal address to the department name, or create / edit '
    'the physical address.',
    'Update department information (ISO rating, station counts) stored as '
    'Codepal user-defined fields (UDFs).',
  ])
  add_para(
    doc,
    'Keeping the master list and related addresses aligned prevents missing '
    'NERIS IDs, wrong department names, and blank or incorrect addresses on '
    'grant documents.'
  )

  add_heading(doc, '2. Who can use this page', 2)
  add_para(
    doc,
    'Only users who are signed in as a Web Admin can open this page. External '
    'users and non-admin staff are redirected to Unauthorized.'
  )
  add_para(
    doc,
    'Navigation: From the admin menu bar, open Admin → Manage FDIDs / '
    'Related Addresses.'
  )
  add_para(
    doc,
    'Page help: Click the help icon (upper right) to expand administrator help '
    'text configured under Manage Help Text (page key: FDIDs (Admin)).'
  )

  add_heading(doc, '3. Fire Department ID list', 2)
  add_para(doc, 'The main grid shows all NERIS IDs from the master table.')
  add_section_table(doc, ['Column', 'Description'], [
    ['Edit', 'View/Edit link — opens the Fire Department ID modal for that row'],
    ['NERIS ID', 'Department NERIS identifier (up to 20 characters; stored uppercase)'],
    ['Fire Department', 'Department name associated with the NERIS ID'],
    ['Inactive', 'When checked on the record, the department is inactive in this list'],
  ])
  add_para(
    doc,
    'The grid supports paging (25 rows per page) and sorting by NERIS ID or '
    'Fire Department (click column headers).'
  )

  add_heading(doc, 'Search and filter', 3)
  add_para(doc, 'Use the toolbar above the grid:')
  add_section_table(doc, ['Control', 'What it does'], [
    ['Search NERIS ID', 'Partial match on NERIS ID (not case-sensitive)'],
    ['Search Fire Department', 'Partial match on department name'],
    ['Hide inactive departments', 'When checked (default), inactive rows are hidden'],
    ['Apply', 'Runs the search/filter and returns to page 1'],
    ['Clear', 'Clears search boxes, re-checks Hide inactive, resets sort to NERIS ID ascending'],
  ])
  add_para(
    doc,
    'Filters apply to the in-memory list loaded for the session; click Apply '
    'after changing search text.'
  )

  add_heading(doc, '4. Add a new NERIS ID', 2)
  add_numbered(doc, [
    'Click Add New NERIS ID.',
    'In the modal, complete the required fields: NERIS ID (required) and '
    'Department Name (required).',
    'Optionally configure fire department address and department information '
    '(see sections 6 and 7).',
    'Set Inactive if the department should not appear when Hide inactive '
    'departments is checked.',
    'Click Save NERIS ID.',
  ])
  add_para(doc, 'On success, the page reloads with a green message: NERIS ID saved successfully.')
  add_para(doc, 'Validation:', bold=True)
  add_bullets(doc, [
    'NERIS ID and Department Name cannot be blank.',
    'NERIS ID must be unique; duplicates show an error in the modal.',
    'NERIS ID is normalized to uppercase as you type.',
  ])

  add_heading(doc, '5. View or edit an existing NERIS ID', 2)
  add_numbered(doc, [
    'In the grid, click View/Edit (or View/Edit followed by the NERIS ID).',
    'The modal opens with NERIS ID, department name, and inactive flag filled in.',
    'If address sync is enabled (see section 8), the system loads matching '
    'addresses and, when possible, the address already linked to that department name.',
    'Make changes and click Save NERIS ID.',
  ])
  add_para(
    doc,
    'Changing the NERIS ID: You may edit the NERIS ID field. Saving replaces '
    'the old ID with the new one in the master list. The new ID must not already exist.'
  )
  add_para(
    doc,
    'Changing the department name: Update Department Name and save. If you link '
    'or create an address (section 6), the linked Codepal address Address Code '
    'is updated to match the new department name.'
  )

  add_heading(doc, '6. Fire department address (link vs create / edit)', 2)
  add_para(
    doc,
    'When address sync is enabled, the modal includes a Fire department address '
    'section below the department name.'
  )
  add_para(
    doc,
    'Link an existing Codepal fire department address or create a new one. Use '
    'Full Address in the dropdown to distinguish departments with the same name.',
    italic=True
  )
  add_para(doc, 'Choose one of two actions:')

  add_heading(doc, 'Option A — Link existing address', 3)
  add_numbered(doc, [
    'Select Link existing address (default).',
    'Open Link to address and choose a row from the dropdown.',
  ])
  add_para(doc, 'Each option is shown as:')
  add_para(
    doc,
    '{Department code} — {Full street address} (Apps: {count}, Users: {count})'
  )
  add_bullets(doc, [
    'Apps — number of Fire Grant applications tied to that address',
    'Users — number of active user/address party links',
  ])
  add_para(
    doc,
    'Use the full address and counts to pick the correct record when several '
    'departments share a similar name (for example, multiple Clovis Fire '
    'Department entries).'
  )
  add_numbered(doc, [
    'When you select an address, the form may switch to Create / Edit Address '
    'mode with that address loaded so you can review or adjust fields before saving.',
    'Click Save NERIS ID.',
  ])
  add_para(
    doc,
    'What “link” does: Updates the selected Codepal address so its Address Code '
    'matches the Department Name you entered. It does not create a new address row.'
  )
  add_para(
    doc,
    'Tip: If no suitable address appears, pick — Create / Edit Address (new) — '
    'at the bottom of the dropdown to switch to create mode.'
  )

  add_heading(doc, 'Option B — Create / Edit Address', 3)
  add_numbered(doc, [
    'Select Create / Edit Address, or choose — Create / Edit Address (new) — '
    'from the link dropdown.',
    'Complete the address fields (see table below).',
    'Click Save NERIS ID.',
  ])
  add_section_table(doc, ['Field', 'Required', 'Notes'], [
    ['Address type', 'Yes', 'Defaults to FS Fire Department when available'],
    ['Street number', 'No', ''],
    ['Direction', 'No', 'e.g. N, S, E, W'],
    ['Street name', 'No', ''],
    ['Suffix', 'No', 'e.g. St, Ave, Blvd'],
    ['City', 'Yes', ''],
    ['State', 'Yes', 'Defaults to New Mexico when configured'],
    ['County', 'Yes', ''],
    ['Zip', 'Yes', 'e.g. 88101; system may create zip if missing'],
  ])
  add_para(
    doc,
    'A full address is required for invoices and legal documents—provide '
    'complete city, state, county, and zip at minimum.'
  )
  add_para(
    doc,
    'What “create” does: Inserts a new Codepal fire department address with the '
    'physical fields you entered and Address Code set to the department name.'
  )
  add_para(
    doc,
    'What “edit” does: If an existing address was loaded (from View/Edit or from '
    'the link dropdown), saving updates that address record instead of creating '
    'a new one.'
  )

  add_heading(doc, '7. Department information (Codepal UDFs)', 2)
  add_para(
    doc,
    'When address sync is enabled, the modal includes Department information '
    '(Codepal UDFs):'
  )
  add_section_table(doc, ['Field', 'Description'], [
    ['ISO Rating', 'ISO classification for the department'],
    ['Main Stations', 'Count of main stations'],
    ['Substations', 'Count of substations'],
    ['Admin Buildings', 'Count of administrative buildings'],
  ])
  add_para(
    doc,
    'These values are stored on the linked or created address and are the same '
    'fields shown on General Information for applicants.'
  )
  add_para(
    doc,
    'Validation: Each count field must be a non-negative whole number (blank is '
    'treated as zero). Invalid values show an error in the modal.'
  )
  add_para(
    doc,
    'UDFs are saved when an address is successfully linked, created, or updated. '
    'If you save a NERIS ID without linking or creating an address, UDF values '
    'are not written.'
  )

  add_heading(doc, '8. Inactive departments', 2)
  add_para(
    doc,
    'Check Inactive in the modal to mark a NERIS ID as inactive in the master list.'
  )
  add_bullets(doc, [
    'Inactive rows show Inactive = True in the grid.',
    'With Hide inactive departments checked (default), inactive rows are hidden '
    'until you clear filters or uncheck the box and click Apply.',
  ])
  add_para(
    doc,
    'Use inactive for departments that should remain in history but should not '
    'appear in normal admin workflows.'
  )

  add_heading(doc, '9. Modal tips', 2)
  add_bullets(doc, [
    'Drag the modal by its title bar to move it on screen.',
    'Close (footer or ×) dismisses the modal without saving.',
    'Errors appear in a red alert at the top of the modal; fix the issue and save again.',
    'After a successful save, the modal closes and the list refreshes.',
  ])

  add_heading(doc, '10. Common errors and what to do', 2)
  add_section_table(doc, ['Message / situation', 'Likely cause', 'What to do'], [
    [
      'NERIS ID cannot be blank',
      'Empty NERIS ID',
      'Enter a valid NERIS ID',
    ],
    [
      'Department Name cannot be blank',
      'Empty name',
      'Enter the official department name',
    ],
    [
      'NERIS ID exists in the list',
      'Duplicate ID',
      'Use a unique NERIS ID or edit the existing row',
    ],
    [
      'Another active address already uses this department name',
      'Duplicate Address Code',
      'Pick a different name, link the correct address, or resolve duplicate in Codepal',
    ],
    [
      'Selected address was not found',
      'Address deleted after dropdown loaded',
      'Close modal, reopen View/Edit, and select again',
    ],
    [
      'City / State / County / Zip / Address type is required…',
      'Missing required address fields',
      'Fill required address fields before saving',
    ],
    [
      '{Field} must be a non-negative whole number',
      'Invalid ISO or station count',
      'Enter 0 or a positive integer',
    ],
    [
      'Unable to resolve or create zip code',
      'Zip not in system',
      'Verify zip format and county',
    ],
    [
      'Multiple addresses match department name (on open)',
      'Duplicate Address Code in Codepal',
      'Use Link to address; choose row using full address and App/User counts',
    ],
  ])
  add_para(
    doc,
    'If address sections are not visible in the modal, address sync may be '
    'turned off in application configuration (EnableFdidAddressSync). Contact '
    'your application administrator or developer; NERIS ID and department name '
    'can still be saved without address linking.'
  )

  add_heading(doc, '11. What this page does not do', 2)
  add_para(doc, 'Understanding these limits avoids unexpected results:')
  add_bullets(doc, [
    'Does not automatically create user accounts or link users to addresses',
    'Does not update applicant General Information department name text already '
    'saved on open applications',
    'Does not create remittance address types used for payment routing',
    'Does not bulk-sync every department in one action—you work one NERIS ID at a time',
  ])

  add_heading(doc, '12. Recommended workflow (NERIS migration cleanup)', 2)
  add_para(
    doc,
    'For a department that is missing NERIS prefill or shows the wrong name on reports:'
  )
  add_numbered(doc, [
    'Search the grid by department name or NERIS ID.',
    'Open View/Edit (or Add New NERIS ID if the department is missing from the master list).',
    'Confirm NERIS ID and Department Name match official NERIS records.',
    'Under Fire department address, link the correct Codepal address or create/edit the physical address.',
    'Enter ISO Rating and station counts if known.',
    'Save and confirm the green success message.',
    'Verify in the applicant flow: open General Information for a test application tied to that department.',
  ])

  add_heading(doc, '13. Quick reference', 2)
  add_section_table(doc, ['Task', 'Steps'], [
    ['Find a department', 'Search NERIS ID or Fire Department → Apply'],
    ['Add NERIS ID', 'Add New NERIS ID → fill form → Save NERIS ID'],
    ['Edit NERIS ID', 'View/Edit → change fields → Save NERIS ID'],
    ['Link to Codepal address', 'Link existing address → choose dropdown row → Save NERIS ID'],
    ['New physical address', 'Create / Edit Address → fill address → Save NERIS ID'],
    ['Hide inactive', 'Keep Hide inactive departments checked → Apply'],
    ['Dismiss modal without saving', 'Close in the modal footer (or ×)'],
    ['Leave page', 'Use the admin menu (e.g. Current Apps) or browser back'],
  ])

  add_heading(doc, 'Document control', 2)
  add_para(
    doc,
    'Related developer docs: fdid-modal-address-sync-implementation-plan.md, '
    'fdid-modal-address-edit-udf-implementation-plan.md, '
    'manage-fdid-list-filter-sort-implementation-plan.md'
  )

  set_update_fields_on_open(doc)
  return doc


if __name__ == '__main__':
  import os

  output = (
    r'i:\NMSFM REPOS\Fire Grant App - CVE Fixes\NMSFM_FGF_CVE\NMSFMFireGrantWF'
    r'\docs\how-to\Manage-FDIDs-Related-Addresses-Admin-How-To.docx'
  )
  temp_output = output + '.generating'
  build_document().save(temp_output)
  try:
    os.replace(temp_output, output)
    print(f'Created: {output}')
  except PermissionError:
    fallback = output.replace('.docx', '-updated.docx')
    if os.path.exists(fallback):
      os.remove(fallback)
    os.replace(temp_output, fallback)
    print(
      f'Original file is open — saved instead as:\n{fallback}\n'
      f'Close the open document and rename, or re-run this script.'
    )
