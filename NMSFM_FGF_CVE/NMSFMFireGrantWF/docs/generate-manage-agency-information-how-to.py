"""Generate Manage Agency Information Admin How-To Word document."""

from datetime import date
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_update_fields_on_open(doc):
  """Prompt Word to refresh fields (including TOC) when the document opens."""
  settings = doc.settings.element
  existing = settings.find(qn('w:updateFields'))
  if existing is None:
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)
  else:
    existing.set(qn('w:val'), 'true')


def add_table_of_contents(doc, levels='2-3'):
  """Insert a Word TOC field for Heading 2–3."""
  toc_title = doc.add_paragraph()
  toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  toc_run = toc_title.add_run('Table of Contents')
  toc_run.bold = True
  toc_run.font.size = Pt(14)

  paragraph = doc.add_paragraph()
  run = paragraph.add_run()

  fld_begin = OxmlElement('w:fldChar')
  fld_begin.set(qn('w:fldCharType'), 'begin')

  instr = OxmlElement('w:instrText')
  instr.set(qn('xml:space'), 'preserve')
  instr.text = f' TOC \\o "{levels}" \\h \\z \\u '

  fld_sep = OxmlElement('w:fldChar')
  fld_sep.set(qn('w:fldCharType'), 'separate')

  run._r.append(fld_begin)
  run._r.append(instr)
  run._r.append(fld_sep)

  paragraph.add_run(
    'Open in Microsoft Word to populate the table of contents.'
  )

  run_end = paragraph.add_run()
  fld_end = OxmlElement('w:fldChar')
  fld_end.set(qn('w:fldCharType'), 'end')
  run_end._r.append(fld_end)

  doc.add_paragraph('')
  page_break = doc.add_paragraph()
  page_break.add_run().add_break(WD_BREAK.PAGE)


def set_cell_shading(cell, fill_hex):
  shading = OxmlElement('w:shd')
  shading.set(qn('w:fill'), fill_hex)
  cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
  doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
  p = doc.add_paragraph()
  run = p.add_run(text)
  if bold:
    run.bold = True
  if italic:
    run.italic = True
  return p


def add_bullets(doc, items):
  for item in items:
    doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc, items):
  for item in items:
    doc.add_paragraph(item, style='List Number')


def add_section_table(doc, headers, rows):
  table = doc.add_table(rows=1, cols=len(headers))
  table.style = 'Table Grid'
  hdr = table.rows[0].cells
  for i, h in enumerate(headers):
    hdr[i].text = h
    set_cell_shading(hdr[i], 'D9E2F3')
    for p in hdr[i].paragraphs:
      for run in p.runs:
        run.bold = True
  for row in rows:
    cells = table.add_row().cells
    for i, val in enumerate(row):
      cells[i].text = val
  doc.add_paragraph('')


def build_document():
  doc = Document()
  style = doc.styles['Normal']
  style.font.name = 'Calibri'
  style.font.size = Pt(11)

  title = doc.add_paragraph()
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  t_run = title.add_run(
    'NMSFM Fire Grant Application\n'
    'Manage Agency Information\nAdmin How-To'
  )
  t_run.bold = True
  t_run.font.size = Pt(18)

  sub = doc.add_paragraph()
  sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
  sub_run = sub.add_run(
    f'For Web Administrators — {date.today().strftime("%B %d, %Y")}'
  )
  sub_run.font.size = Pt(12)

  doc.add_paragraph('')
  add_para(doc, 'Project: NMSFM Fire Grant Web Application')
  add_para(doc, 'Document version: 1.0')
  add_para(doc, 'Audience: Fire Grant web administrators')
  add_para(
    doc,
    'Page: Admin → Manage Agency Information (/Admin/ManageAgencyInformation)'
  )
  add_para(
    doc,
    'Source markdown: docs/how-to/manage-agency-information-how-to.md'
  )
  doc.add_paragraph('')

  add_table_of_contents(doc)

  add_heading(doc, '1. What this page is for', 2)
  add_para(
    doc,
    'The Manage Agency Information page lets web administrators view and edit '
    'the agency record associated with their login session. This is the same '
    'agency contact data maintained in the LegacyApp Agency form, now available '
    'in the Fire Grant web application.'
  )
  add_para(doc, 'Agency information is used across the application—for example:')
  add_bullets(doc, [
    'Grant reports — agency name, address, and report image on printed documents',
    'Support menu — technical and fire services support email recipients',
    'Agency-level settings — user-defined fields (UDFs) configured in Codepal',
  ])
  add_para(
    doc,
    'The page opens an Agency Information modal automatically when you navigate '
    'to it. You edit fields across three tabs and click Save to persist changes '
    'to the Codepal database.'
  )

  add_heading(doc, '2. Who can use this page', 2)
  add_para(
    doc,
    'Only users who are signed in as a Web Admin can open this page. External '
    'users and non-admin staff are redirected to Unauthorized.'
  )
  add_para(
    doc,
    'Navigation: From the admin menu bar, open Admin → Manage Agency Information.'
  )
  add_para(
    doc,
    'The menu item is also available from the application sidebar when editing '
    'a grant application.'
  )
  add_para(
    doc,
    'Page help: Click the help icon (upper right) for administrator help text '
    '(page key: Manage Agency Information (Admin)).'
  )
  add_para(
    doc,
    'Agency scope: You can edit only the agency tied to your login session. '
    'This page does not list or switch between multiple agencies.'
  )

  add_heading(doc, '3. Opening the Agency Information modal', 2)
  add_numbered(doc, [
    'Open Admin → Manage Agency Information.',
    'The Agency Information modal opens automatically with current data loaded.',
    'Use the General, Advanced, and Support Email tabs to review or edit fields.',
    'Click Save to persist changes, or Close to return to Current Apps without saving.',
  ])
  add_para(
    doc,
    'After a successful save, a green success message appears above the page '
    'title and the modal reopens with updated values.'
  )
  add_para(
    doc,
    'If validation fails, a red alert appears at the top of the modal and the '
    'modal stays open so you can correct the issue.'
  )

  add_heading(doc, '4. Inactive checkbox', 2)
  add_para(
    doc,
    'At the top right of the modal, the Inactive checkbox marks the agency '
    'record as inactive in Codepal (ExternalId: 0 = active, 1 = inactive).'
  )
  add_para(
    doc,
    'Use this only when directed by your agency administrator. In most day-to-day '
    'Fire Grant administration, leave Inactive unchecked.'
  )

  add_heading(doc, '5. General tab', 2)
  add_para(
    doc,
    'The General tab contains agency contact information and the report image.'
  )

  add_heading(doc, 'Contact fields', 3)
  add_section_table(doc, ['Field', 'Max length', 'Notes'], [
    ['Name', '50', 'Official agency name (AgencyName)'],
    ['Sub Name', '50', 'Secondary name or subtitle (AgencySubName)'],
    ['Address', '50', 'Street address'],
    ['City', '50', ''],
    ['State / Zip', '—', 'State dropdown plus zip text (zip max 20 characters)'],
    ['Country', '—', 'Country dropdown'],
    ['Phone', '25', ''],
    ['Fax', '25', ''],
    ['E-mail', '100', 'Validated as an email address format'],
  ])
  add_para(
    doc,
    'Layout: contact fields appear in the left column; the report image panel '
    'is on the right.'
  )

  add_heading(doc, 'Report image', 3)
  add_para(
    doc,
    'The Report Image appears on printed grant reports. It is stored in the '
    'agency record as a binary image.'
  )
  add_section_table(doc, ['Action', 'How'], [
    ['View current image', '150×150 thumbnail preview when an image is saved'],
    ['Upload new image', 'Use Choose File, then click Save'],
    ['Clear image', 'Click Clear Report Image, then click Save'],
  ])
  add_para(doc, 'Allowed file types: .bmp, .jpg, .jpeg, .gif, .png')
  add_para(
    doc,
    'Tip: Click Clear Report Image before Save to remove the image from the '
    'database. Clearing only hides the preview until you save.'
  )

  add_heading(doc, 'Created and Last Updated', 3)
  add_bullets(doc, [
    'Created — date and time the agency record was first inserted',
    'Last Updated — date and time of the most recent save',
  ])

  add_heading(doc, '6. Advanced tab', 2)
  add_para(
    doc,
    'The Advanced tab shows agency-level user-defined fields (UDFs) configured '
    'in Codepal. Fields are grouped under category headings.'
  )

  add_heading(doc, 'Field types you may see', 3)
  add_section_table(doc, ['Control type', 'How it works'], [
    ['Text box', 'Free-text entry (default for most field types)'],
    ['Drop-down list', 'Choose one option from a predefined list'],
    ['Check box list', 'Select one or more options from a predefined list'],
  ])

  add_heading(doc, 'Empty state', 3)
  add_para(
    doc,
    'If no agency UDF categories are configured, the tab displays: '
    '"There are no additional fields defined for this record."'
  )

  add_heading(doc, 'Required fields', 3)
  add_para(
    doc,
    'Some UDF fields may be marked required in Codepal. If a required field is '
    'blank when you click Save, the modal shows an error such as: '
    'Required field \'{field name}\' must be completed.'
  )

  add_heading(doc, '7. Support Email tab', 2)
  add_para(
    doc,
    'The Support Email tab configures who receives email when users submit the '
    'Technical Support and Fire Services Support forms from the Support menu.'
  )
  add_section_table(doc, ['Field', 'Purpose'], [
    ['Technical Support Email', 'Recipients for technical / IT support requests'],
    ['Fire Services Support Email', 'Recipients for fire services program questions'],
  ])

  add_heading(doc, 'Format', 3)
  add_bullets(doc, [
    'Enter one or more email addresses.',
    'Separate multiple addresses with semicolons (;).',
    'Example: admin@example.gov; helpdesk@example.gov',
  ])

  add_heading(doc, 'Blank fields', 3)
  add_para(
    doc,
    'Leave a field blank to use the web.config fallback address configured by '
    'your application deploy team.'
  )

  add_heading(doc, 'Validation', 3)
  add_para(
    doc,
    'Each semicolon-separated address is validated when you save. Invalid '
    'addresses show an error in the modal.'
  )

  add_heading(doc, '8. Saving changes', 2)
  add_numbered(doc, [
    'Edit fields on any tab (General, Advanced, or Support Email).',
    'Click Save in the modal footer.',
    'On success: green message "Agency information has been saved." and modal reopens.',
    'On error: red alert in modal; fix the issue and click Save again.',
  ])
  add_para(
    doc,
    'Save writes all tabs in one action — agency contact fields, report image, '
    'UDF values, and support email settings are persisted together.'
  )

  add_heading(doc, '9. Closing without saving', 2)
  add_para(
    doc,
    'Click Close in the modal footer to navigate to Current Apps (/Admin/Home). '
    'Close does not save changes made since the last successful save.'
  )
  add_para(
    doc,
    'You can also dismiss the modal with the × button in the header; use Close '
    'to return to the admin home page.'
  )

  add_heading(doc, '10. Common errors and what to do', 2)
  add_section_table(doc, ['Message / situation', 'Likely cause', 'What to do'], [
    [
      'Agency session is missing. Please log in again.',
      'Session expired',
      'Sign out and sign in again',
    ],
    [
      'Agency record was not found.',
      'Database record missing',
      'Contact your application administrator',
    ],
    [
      'Report image must be a .bmp, .jpg, .jpeg, .gif, or .png file.',
      'Unsupported upload type',
      'Choose a supported image format',
    ],
    [
      'Required field \'…\' must be completed.',
      'Blank required UDF',
      'Fill in the named field on Advanced tab',
    ],
    [
      'Technical Support Email: invalid email \'…\'',
      'Malformed address',
      'Correct email; use semicolons between addresses',
    ],
    [
      'Fire Services Support Email: invalid email \'…\'',
      'Malformed address',
      'Correct email; use semicolons between addresses',
    ],
    [
      'Unable to save agency information.',
      'Database or service error',
      'Retry; contact administrator if it persists',
    ],
    [
      'Unable to save support email settings.',
      'Settings update failed',
      'Retry; contact administrator if it persists',
    ],
    [
      'Redirect to Unauthorized',
      'Not a web admin',
      'Sign in with an admin account',
    ],
  ])

  add_heading(doc, '11. What this page does not do', 2)
  add_bullets(doc, [
    'Does not provide a list of agencies or let you switch agencies',
    'Does not include the LegacyApp Certifications tab',
    'Does not show audit history or change logs',
    'Does not manage county fields (omitted from the web form)',
    'Does not update fire department (FDID) records — use Manage FDIDs / Related Addresses',
  ])

  add_heading(doc, '12. Recommended workflows', 2)

  add_heading(doc, 'Update agency contact information for reports', 3)
  add_numbered(doc, [
    'Open Admin → Manage Agency Information.',
    'On the General tab, update Name, address, phone, and E-mail.',
    'Click Save.',
    'Open a sample grant report and confirm the agency block is correct.',
  ])

  add_heading(doc, 'Replace the report logo/image', 3)
  add_numbered(doc, [
    'Open the modal → General tab.',
    'Under Report Image, choose a .png or .jpg file.',
    'Click Save.',
    'Preview a printed report to confirm the image appears correctly.',
  ])

  add_heading(doc, 'Update support email recipients', 3)
  add_numbered(doc, [
    'Open the modal → Support Email tab.',
    'Enter semicolon-separated addresses.',
    'Click Save.',
    'Submit a test message from the Support menu to confirm delivery.',
  ])

  add_heading(doc, 'Configure agency UDFs', 3)
  add_numbered(doc, [
    'Open the modal → Advanced tab.',
    'Complete fields under each category heading.',
    'Click Save.',
    'Verify values where agency UDFs are consumed.',
  ])

  add_heading(doc, '13. Quick reference', 2)
  add_section_table(doc, ['Task', 'Steps'], [
    ['Open agency editor', 'Admin → Manage Agency Information'],
    ['Edit contact info', 'General tab → edit fields → Save'],
    ['Upload report image', 'General tab → Choose File → Save'],
    ['Remove report image', 'General tab → Clear Report Image → Save'],
    ['Edit UDFs', 'Advanced tab → edit fields → Save'],
    ['Set support recipients', 'Support Email tab → enter emails → Save'],
    ['Return to admin home', 'Close'],
    ['Get help', 'Click help icon (upper right)'],
  ])

  add_heading(doc, 'Document control', 2)
  add_para(
    doc,
    'Related developer docs: manage-agency-information-implementation-plan.md, '
    'support-email-admin-tab-implementation-plan.md'
  )

  set_update_fields_on_open(doc)
  return doc


if __name__ == '__main__':
  output = (
    r'i:\NMSFM REPOS\Fire Grant App - CVE Fixes\NMSFM_FGF_CVE\NMSFMFireGrantWF'
    r'\docs\how-to\Manage-Agency-Information-Admin-How-To.docx'
  )
  temp_output = output + '.generating'
  build_document().save(temp_output)
  try:
    os.replace(temp_output, output)
    print(f'Created: {output}')
  except PermissionError:
    fallback = output.replace('.docx', '-updated.docx')
    if os.path.exists(fallback):
      os.remove(fallback)
    os.replace(temp_output, fallback)
    print(
      f'Original file is open — saved instead as:\n{fallback}\n'
      f'Close the open document and rename, or re-run this script.'
    )
